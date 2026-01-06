import requests, json, base64, win32com.client, os, time
from docx import Document
from decouple import config

api_key = config('GEMINI_API_KEY', default='your_api_key_here')
url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent?key={api_key}"

def convert_doc_to_docx_windows(doc_path):
    word = win32com.client.Dispatch("Word.Application")
    doc = word.Documents.Open(doc_path)
    docx_path = doc_path + "x"
    doc.SaveAs(docx_path, FileFormat=16)  # 16 = wdFormatDocumentDefault (docx)
    doc.Close()
    word.Quit()
    return docx_path

def read_docx_with_tables(file):
    doc = Document(file)
    full_text = []

    # Текст из параграфов
    for para in doc.paragraphs:
        full_text.append(para.text)

    # Таблицы
    for table in doc.tables:
        full_text.append("\n[Таблица]")
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            full_text.append(" | ".join(cells))  # можно заменить на CSV-стиль
        full_text.append("[Конец таблицы]\n")

    return "\n".join(full_text)

def chatgpt(prompt):
    # Задержка для избежания лимита запросов (429)
    time.sleep(2)  # 2 секунды задержки
    
    headers = {
        "Content-Type": "application/json"
    }

    payload = {
        "contents": [
            {
                "parts": [
                    {"text": prompt+", ответ выводить без стилизации"}
                ]
            }
        ]
    }
    try:
        response = requests.post(url, headers=headers, data=json.dumps(payload))
        if response.status_code != 200:
            return f'<div class="alert alert-danger">Ошибка API: {response.status_code} - {response.text}</div>'
        result = response.json()
        print("API Response:", result)  # Для отладки
        if 'candidates' in result and result['candidates']:
            text = result['candidates'][0]['content']['parts'][0]['text']
            return text
        else:
            return '<div class="alert alert-danger">Ошибка: Неверный ответ от API</div>'
    except requests.exceptions.RequestException as e:
        return f'<div class="alert alert-danger">Ошибка сети: {str(e)}</div>'
    except KeyError as e:
        return f'<div class="alert alert-danger">Ошибка обработки ответа: {str(e)}</div>'
    except Exception as e:
        return f'<div class="alert alert-danger">Неизвестная ошибка: {str(e)}</div>'


def send_file(file_paths, prompt):
    file_contents = []
    for f in file_paths:
        if f.endswith('.docx'):
            with open(f, 'rb') as file:
                file_contents.append(read_docx_with_tables(file))
        else:
            return "Файлҳо бояд формати docx дошта бошанд."
    file_contents = "----начало нового файла----".join(file_contents)
    return chatgpt(prompt + file_contents)