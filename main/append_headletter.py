import tempfile
import os
from docx import Document as DocxDocument
from docxcompose.composer import Composer
from django.core.files.base import ContentFile

from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

def add_styled_text(doc, text):
    """Добавляет текст с готовой стилизацией для печати"""
    for line in text.split('\n'):
        line = line.strip()
        if not line:
            continue
        # Если есть маркеры от ИИ
        if line.startswith("[H]"):
            h = doc.add_heading(line[3:].strip(), level=1)
            run = h.runs[0]
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            run.font.size = Pt(16)
            run.font.bold = True
            h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif line.startswith("[P]") or not line.startswith("[OL]") and not line.startswith("[UL]"):
            p = doc.add_paragraph(line[3:].strip() if line.startswith("[P]") else line)
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                run.font.size = Pt(14)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY


def generate_final_doc(document):
    print(document)
    if document.school.headletter:
        base_doc = DocxDocument(document.school.headletter.path)
    else:
        base_doc = DocxDocument()

    composer = Composer(base_doc)

    if document.file:
        user_doc = DocxDocument(document.file.path)
        composer.append(user_doc)

    if document.content:
        text_doc = DocxDocument()
        add_styled_text(text_doc, document.content)
        composer.append(text_doc)

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        temp_path = tmp.name

    composer.save(temp_path)

    with open(temp_path, "rb") as f:
        document.file.save(f"{document.title}.docx", ContentFile(f.read()), save=True)

    os.remove(temp_path)
    return document
